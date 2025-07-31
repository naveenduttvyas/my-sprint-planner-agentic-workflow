import os
import requests
from pathlib import Path
from dotenv import load_dotenv
from agent_executor.agents.classify_story_type import classify_story_type
from agent_executor.agents.code_updater import update_code_llm

load_dotenv()

def ask_llm(api_key, prompt):
    url = "https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent"
    response = requests.post(url, params={"key": api_key}, json={"contents": [{"parts": [{"text": prompt}]}]})
    response.raise_for_status()
    data = response.json()
    return data.get("candidates", [{}])[0].get("content", {}).get("parts", [{}])[0].get("text", "").strip()

def run(state):
    api_key = os.getenv("GEMINI_API_KEY")
    repo_path = os.getenv("LOCAL_REPO_PATH", "sample-api")  # Local path to scanned repo

    for story in state["ai_stories"]:
        summary = story.get("summary") or ""
        description = story.get("description") or ""
        full_text = summary + "\n" + description

        story_type = classify_story_type(summary, description)
        story["type"] = story_type
        print(f"[DEBUG] Story '{story['key']}' classified as: {story_type}")

        if story_type == "document":
            from docx import Document

            # Ask LLM to generate detailed document content
            doc_prompt = f"""
        Write a technical design or documentation draft for the following user story.
        Structure it with clear headings, bullet points, and paragraphs:

        ---
        {full_text}
        """
            doc_text = ask_llm(api_key, doc_prompt)

            # Create Word document
            repo_root = Path(os.getenv("LOCAL_REPO_PATH", "sample-api")).resolve()
            doc_path = repo_root / "docs" / story["key"]
            doc_path.mkdir(parents=True, exist_ok=True)

            doc_file = doc_path / "story_documentation.docx"
            document = Document()

            # Parse and format document text (simple structure)
            document.add_heading(summary.strip(), level=0)
            for para in doc_text.split("\n\n"):
                para = para.strip()
                if not para:
                    continue
                elif para.startswith("- ") or para.startswith("* "):
                    for bullet in para.splitlines():
                        document.add_paragraph(bullet.strip(" -"), style='List Bullet')
                elif para.endswith(":"):
                    document.add_heading(para, level=1)
                else:
                    document.add_paragraph(para)

            document.save(doc_file)
            story["updated_file_path"] = str(doc_file)
            print(f"[INFO] Document for story {story['key']} saved at {doc_file}")
            continue


        elif story_type == "code":
            # Step 1: Ask LLM for clean functional code
            gen_code_prompt = f"""
Generate only the clean code (function, class, or logic) that addresses the following user story.
Avoid comments or explanations. Only return the exact code block to insert.

---
{full_text}
"""
            generated_code = ask_llm(api_key, gen_code_prompt)
            story["code"] = generated_code

            # Step 2: Use generic LLM updater to locate and patch the right file
            updated_path = update_code_llm(
                repo_root=repo_path,
                story={"summary": summary, "description": description},
                generated_code=generated_code
            )
            story["updated_file_path"] = updated_path

            print(f"[INFO] Code for story {story['key']} written to {updated_path}")

        else:
            story["code"] = "# Unknown story type. Skipped."

    return state